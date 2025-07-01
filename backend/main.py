from fastapi import FastAPI, HTTPException, Depends
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from docx import Document
import os
from datetime import datetime
from sqlalchemy import create_engine, Column, Integer, String, Date
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker, Session
from dotenv import load_dotenv
import subprocess

# Load environment variables
load_dotenv()

# Initialize FastAPI app
app = FastAPI()

# CORS setup
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Database setup
SQLALCHEMY_DATABASE_URL = os.getenv("DATABASE_URL")
engine = create_engine(SQLALCHEMY_DATABASE_URL)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()

# DB Models
class OfferLetter(Base):
    __tablename__ = "offer_letters"
    id = Column(Integer, primary_key=True, index=True)
    name = Column(String, index=True)
    duration = Column(String)
    start_date = Column(String)
    end_date = Column(String)
    generated_on = Column(Date)

# Pydantic models
class OfferLetterRequest(BaseModel):
    offer_letter_id: int

class OfferLetterCreate(BaseModel):
    name: str
    duration: str
    start_date: str  # Accept as string, parse later
    end_date: str

# Helper Functions
def convert_to_datetime(date_string: str) -> datetime:
    try:
        return datetime.strptime(date_string, "%B %d, %Y")
    except ValueError:
        return datetime.strptime(date_string, "%Y-%m-%d")

def remove_page_breaks_and_adjust_spacing(doc: Document):
    for para in doc.paragraphs:
        # Remove manual page breaks
        if para.text == "":
            continue
        for run in para.runs:
            if run.text == "\n":
                run.text = ""  # Remove page breaks

        # Adjust paragraph spacing
        para.paragraph_format.space_after = 0  # Remove extra space after paragraphs
        para.paragraph_format.space_before = 0  # Remove extra space before paragraphs

    # Loop through headers and footers to remove page breaks as well
    for section in doc.sections:
        for header_footer in [section.header, section.footer]:
            for para in header_footer.paragraphs:
                for run in para.runs:
                    if run.text == "\n":
                        run.text = ""  # Remove page breaks
                para.paragraph_format.space_after = 0
                para.paragraph_format.space_before = 0

def clear_header(doc: Document):
    for section in doc.sections:
        for header in section.header.paragraphs:
            header.clear()

def replace_placeholders(doc: Document, replacements: dict):
    def replace_in_paragraph(paragraph):
        full_text = ''.join(run.text for run in paragraph.runs)
        replaced = False
        for key, val in replacements.items():
            if key in full_text:
                full_text = full_text.replace(key, val)
                replaced = True
        if replaced:
            for run in paragraph.runs:
                run.text = ''
            if paragraph.runs:
                paragraph.runs[0].text = full_text


    for para in doc.paragraphs:
        replace_in_paragraph(para)

    # Clear header content
    clear_header(doc)

    # Replace in the rest of the document sections
    for section in doc.sections:
        for header_footer in [section.header, section.footer]:
            for para in header_footer.paragraphs:
                replace_in_paragraph(para)

def generate_word_document(name: str, duration: str, start_date: str, end_date: str) -> str:
    start_date = convert_to_datetime(str(start_date))
    end_date = convert_to_datetime(str(end_date))

    start_date_str = start_date.strftime("%B %d, %Y")
    end_date_str = end_date.strftime("%B %d, %Y")

    current_date_str = datetime.today().strftime("%B %d, %Y")  # Get today's date

    replacements = {
        "<name>": name,
        "<duration>": duration,
        "<start_date>": start_date_str,
        "<end_date>": end_date_str,
        "<current_date>": current_date_str  # Add this line
    }

    doc_path = os.path.join(os.path.dirname(__file__), "template_offer_letter.docx")
    doc = Document(doc_path)

    replace_placeholders(doc, replacements)

    # Clear the header content before saving
    clear_header(doc)

    output_docx = f"updated_offer_letter_{name.replace(' ', '_')}.docx"
    doc.save(output_docx)
    return output_docx

# DB Dependency
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# Routes
@app.post("/create_offer_letter/")
def create_offer_letter(offer: OfferLetterCreate, db: Session = Depends(get_db)):
    try:
        # Convert incoming strings to date objects first
        start_date_obj = datetime.strptime(offer.start_date, "%m-%d-%Y")
        end_date_obj = datetime.strptime(offer.end_date, "%m-%d-%Y")

        # Now format into "January 1, 2025" format
        start_date_str = start_date_obj.strftime("%B %d, %Y")
        end_date_str = end_date_obj.strftime("%B %d, %Y")

        new_offer = OfferLetter(
            name=offer.name,
            duration=offer.duration,
            start_date=start_date_str,
            end_date=end_date_str
        )

        db.add(new_offer)
        db.commit()
        db.refresh(new_offer)
        return {"id": new_offer.id}
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error: {e}")

@app.post("/generate_offer_letter/")
async def generate_offer_letter(request: OfferLetterRequest, db: Session = Depends(get_db)):
    offer_letter = db.query(OfferLetter).filter(OfferLetter.id == request.offer_letter_id).first()
    if not offer_letter:
        raise HTTPException(status_code=404, detail="Offer letter not found")

    word_file = generate_word_document(
        offer_letter.name,
        offer_letter.duration,
        offer_letter.start_date,
        offer_letter.end_date
    )

    return FileResponse(
        word_file,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=os.path.basename(word_file),
        headers={"Content-Disposition": f"inline; filename={os.path.basename(word_file)}"},
    )
@app.get("/get_offer_letter_names/")
def get_offer_letter_names(db: Session = Depends(get_db)):
    offer_letters = db.query(OfferLetter.id, OfferLetter.name).all()
    if not offer_letters:
        raise HTTPException(status_code=404, detail="No offer letters found")
    return [{"id": ol.id, "name": ol.name} for ol in offer_letters]

@app.get("/offer_letters/")
def get_offer_letters(db: Session = Depends(get_db)):
    offer_letters = db.query(OfferLetter).all()
    if not offer_letters:
        raise HTTPException(status_code=404, detail="No offer letters found")
    return offer_letters
